# -*- coding: utf-8 -*-
"""
Genere la version Word de la charte.

Ce n'est pas une reproduction du PDF, et cela ne peut pas l'etre : les fonds perdus,
les planches de composants et les aplats bord a bord n'existent pas dans Word. C'est
la version EDITABLE du meme fond, composee dans l'idiome de Word : styles de titre,
tableaux a filets horizontaux, nuancier en cellules teintees.

Le PDF reste la reference visuelle. Ce fichier sert a relire, annoter et amender.

    python generer-docx.py
"""
import io, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SORTIE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      "..", "AI5D_Digital_Design_System_Charte.docx")

ENCRE   = RGBColor(0x05, 0x1C, 0x2C)
ACTION  = RGBColor(0x22, 0x51, 0xFF)
TEXTE   = RGBColor(0x2B, 0x3A, 0x45)
FAIBLE  = RGBColor(0x61, 0x6F, 0x78)
ERREUR  = RGBColor(0xB4, 0x23, 0x18)

SERIF = "Fraunces"      # Georgia en substitution si la police n'est pas installee
SANS  = "Inter"         # Segoe UI en substitution
MONO  = "JetBrains Mono"

doc = Document()


# --------------------------------------------------------------------------- outils
def police(run, nom, taille, couleur=TEXTE, gras=False, italique=False):
    run.font.name = nom
    run.font.size = Pt(taille)
    run.font.color.rgb = couleur
    run.font.bold = gras
    run.font.italic = italique
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for att in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(att), nom)
    return run


def para(espace_avant=0, espace_apres=6, interligne=1.35, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(espace_avant)
    pf.space_after = Pt(espace_apres)
    pf.line_spacing = interligne
    if align is not None:
        p.alignment = align
    return p


def corps(texte, taille=9.5, couleur=TEXTE, avant=0, apres=7):
    """Un paragraphe ou **gras** est interprete."""
    p = para(avant, apres)
    for i, morceau in enumerate(texte.split("**")):
        if not morceau:
            continue
        police(p.add_run(morceau), SANS, taille,
               ENCRE if i % 2 else couleur, gras=bool(i % 2))
    return p


def surtitre(texte):
    p = para(14, 2, 1.0)
    police(p.add_run(texte.upper()), MONO, 7.5, ACTION, gras=True)
    return p


def titre_chapitre(numero, titre):
    p = para(0, 3, 1.0)
    police(p.add_run("%02d" % numero), MONO, 9, ACTION, gras=True)
    p2 = para(0, 10, 1.1)
    police(p2.add_run(titre), SERIF, 21, ENCRE)
    return p2


def etiquette(texte):
    p = para(15, 4, 1.0)
    police(p.add_run(texte.upper()), MONO, 7, FAIBLE, gras=True)
    filet(p)
    return p


def filet(p, position="top"):
    pPr = p._element.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    el = OxmlElement('w:' + position)
    el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
    el.set(qn('w:space'), '6');    el.set(qn('w:color'), 'E7E0D6')
    bd.append(el); pPr.append(bd)


def teinter(cellule, hexa):
    tcPr = cellule._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexa.lstrip('#'))
    tcPr.append(sh)


def filets_horizontaux(table):
    """Filets horizontaux seulement : jamais de grille fermee."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for bord, couleur, taille in (('top', 'FFFFFF', '0'), ('left', 'FFFFFF', '0'),
                                  ('bottom', 'FFFFFF', '0'), ('right', 'FFFFFF', '0'),
                                  ('insideH', 'E7E0D6', '4'), ('insideV', 'FFFFFF', '0')):
        el = OxmlElement('w:' + bord)
        el.set(qn('w:val'), 'single' if taille != '0' else 'none')
        el.set(qn('w:sz'), taille); el.set(qn('w:space'), '0')
        el.set(qn('w:color'), couleur)
        borders.append(el)
    tblPr.append(borders)


def tableau(entetes, lignes, largeurs=None, gras_col0=True, nuancier=False):
    """nuancier : la premiere colonne porte un hex, rendu en aplat."""
    t = doc.add_table(rows=1, cols=len(entetes))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    filets_horizontaux(t)
    for i, e in enumerate(entetes):
        c = t.rows[0].cells[i]
        c.paragraphs[0].paragraph_format.space_after = Pt(4)
        police(c.paragraphs[0].add_run(e.upper()), MONO, 6.5, FAIBLE, gras=True)
    for ligne in lignes:
        cellules = t.add_row().cells
        for i, val in enumerate(ligne):
            p = cellules[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            if nuancier and i == 0:
                teinter(cellules[i], val)
                police(p.add_run(" "), SANS, 8)
                continue
            mono = val.startswith("#") or val.startswith("--") or val.startswith("[")
            couleur = ERREUR if val.endswith("✖") else (ENCRE if i == 0 else TEXTE)
            texte = val.rstrip("✖")
            police(p.add_run(texte), MONO if mono else SANS, 8 if mono else 8.5,
                   couleur, gras=(i == 0 and gras_col0 and not mono))
    if largeurs:
        for r in t.rows:
            for i, l in enumerate(largeurs):
                r.cells[i].width = Cm(l)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def encadre(texte):
    p = para(10, 10, 1.4)
    pPr = p._element.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), 'F4EFE7')
    pPr.append(sh)
    bd = OxmlElement('w:pBdr'); el = OxmlElement('w:left')
    el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '18')
    el.set(qn('w:space'), '8'); el.set(qn('w:color'), '2251FF')
    bd.append(el); pPr.append(bd)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    for i, morceau in enumerate(texte.split("**")):
        if morceau:
            police(p.add_run(morceau), SANS, 9, ENCRE if i % 2 else TEXTE, gras=bool(i % 2))
    return p


def code(texte):
    p = para(8, 10, 1.3)
    pPr = p._element.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), 'F4EFE7')
    pPr.append(sh)
    p.paragraph_format.left_indent = Cm(0.3)
    lignes = texte.split("\n")
    for i, l in enumerate(lignes):
        if i:
            p.add_run().add_break()
        police(p.add_run(l), MONO, 8, ENCRE)
    return p


def saut():
    doc.add_page_break()


# --------------------------------------------------------------------------- page
for s in doc.sections:
    s.page_height, s.page_width = Cm(29.7), Cm(21.0)
    s.top_margin = Cm(2.4); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)

normal = doc.styles['Normal']
normal.font.name = SANS
normal.font.size = Pt(9.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), SANS)


# --------------------------------------------------------------------------- titre
p = para(150, 4, 1.0)
police(p.add_run("CHARTE DE RÉFÉRENCE"), MONO, 9, ACTION, gras=True)
p = para(0, 0, 1.0)
police(p.add_run("AI5D"), SERIF, 40, ENCRE)
p = para(0, 12, 1.0)
police(p.add_run("Digital Design System"), SERIF, 40, ACTION)
p = para(0, 6, 1.2)
police(p.add_run("Un seul ADN. Quatre densités."), SERIF, 15, ENCRE)
corps("Le registre applicatif de l’écosystème AI5D. Ce que tout produit hérite. Les cinq "
      "libertés qu’il obtient. Et les trois règles qu’aucune machine ne le laissera "
      "enfreindre.", 9.5, FAIBLE, apres=30)
p = para(0, 0, 1.0)
filet(p)
police(p.add_run("@ai5d/design-system · v0.1.1 · Édition 2026"), MONO, 7.5, FAIBLE, gras=True)

encadre("**Sur cette version.** Le PDF compagnon est la référence visuelle : ses planches, "
        "ses aplats et ses reproductions d’écran n’existent pas dans Word. Ce fichier porte "
        "le même fond sous une forme éditable, pour relire, annoter et amender. En cas "
        "d’écart de mise en page entre les deux, c’est le PDF qui fait foi ; en cas d’écart "
        "de valeur, c’est noyau/jetons.css.")
saut()


# --------------------------------------------------------------------------- sommaire
surtitre("Le document")
p = para(0, 10, 1.1)
police(p.add_run("Seize chapitres. Une seule question : qu’est-ce qui se copie."), SERIF, 21, ENCRE)
corps("Deux chartes dérivées en six mois. Quatre jetons qui ont divergé sans que personne "
      "ne l’ait décidé. Ce document existe pour qu’il n’y ait pas de troisième dérivation : "
      "il décrit ce qu’un produit AI5D reçoit le jour où il est créé.")

SOMMAIRE = [
    ("01", "Deux registres, une famille", "Ce que la marque garde, ce qu’elle accorde"),
    ("02", "Les trois couches", "Noyau, densités, écosystème"),
    ("03", "Les jetons de marque", "Six valeurs importées, jamais définies"),
    ("04", "Palette claire", "Papier tiède, encre, sémantiques"),
    ("05", "Palette sombre", "La profondeur vient de la lumière"),
    ("06", "Les écarts mesurés", "Quatre divergences, trois bonnes décisions"),
    ("07", "Typographie", "Fraunces signe, Inter travaille"),
    ("08", "Les densités", "Quatre profils, deux règles"),
    ("09", "Les huit composants", "Ce que chacun garantit"),
    ("10", "Planches", "Les composants, clair et sombre"),
    ("11", "Le gabarit d’authentification", "Une disposition mesurée, pas proposée"),
    ("12", "Mobile d’abord", "Le téléphone est le point de départ"),
    ("13", "Iconographie et mouvement", "Épaisseur 1,75 et trois durées"),
    ("14", "La voix", "Formulations de référence"),
    ("15", "Les gardes", "Ce qu’une machine vérifie à notre place"),
    ("16", "Adoption", "Une ligne, et un versionnement sévère"),
]
tableau(["#", "Chapitre", "Ce qu’il tranche"], SOMMAIRE, [1.4, 6.4, 8.4])
encadre("**Statut.** Source de vérité du registre applicatif AI5D. En cas de contradiction "
        "avec la charte d’un produit, c’est ce document qui juge. Les valeurs exécutables "
        "vivent dans noyau/jetons.css et densites/profils.css ; ce document les explique, "
        "il ne les remplace pas.")
saut()


# --------------------------------------------------------------------------- 01
surtitre("Deux registres, une famille")
titre_chapitre(1, "La marque ne bouge pas. Ce sont les produits qui vivent.")
corps("Une page de présentation se parcourt en dix secondes. Un produit s’habite pendant "
      "vingt minutes. Les deux ne peuvent pas obéir aux mêmes règles.")
corps("Le registre institutionnel garde les angles vifs, l’absence totale d’ombre et Inter "
      "pour seule police. Il reste sous l’autorité de la charte mère. Ce document définit "
      "l’autre moitié, une fois, pour tous les produits.")
etiquette("Les cinq libertés accordées")
tableau(["Liberté", "Doctrine institutionnelle", "Pourquoi le registre applicatif s’en écarte"], [
    ("Coins arrondis 4 · 10 · 16 px", "Rayon 0 partout",
     "Une carte à angles vifs sur un fond de page se lit comme un défaut de rendu, pas comme une intention"),
    ("Trois niveaux d’élévation", "Aucune ombre, jamais",
     "Une élévation pose une carte au-dessus de la page et dit, sans un mot, où se trouve la seule chose à faire"),
    ("Fraunces en serif d’affichage", "Aucune police serif",
     "Toutes les interfaces concurrentes composent en sans géométrique. La serif rend le produit reconnaissable en une demi-seconde"),
    ("Surfaces tièdes #FAF7F2", "Blanc pur",
     "Le blanc pur fatigue au-delà de vingt minutes de lecture, et le produit se lit plus longtemps que la vitrine"),
    ("Mode sombre complet", "Absent de la charte mère",
     "Une partie des sessions se fait la nuit, sur téléphone, dans une pièce peu éclairée"),
], [4.4, 3.6, 8.2])
etiquette("Ce qui reste non négociable")
corps("**Les six jetons de marque.** Encre, navy, bleu signal et ses deux états, blanc. Ils "
      "sont importés depuis la source institutionnelle, et un test relit cette source à "
      "chaque exécution.")
corps("**Le « 5 » incliné.** À moins cinq degrés, en bleu, dans toutes les variantes et sur "
      "tous les fonds. Ne jamais le redresser, ne jamais le recolorer.")
corps("**La différenciation par le nom.** Les branches partagent une seule identité "
      "visuelle. Un produit ne prend jamais sa propre couleur.")
encadre("Ces cinq libertés ne sont pas des tolérances laissées à l’appréciation de chacun : "
        "elles forment une **liste fermée**. Un produit qui a besoin d’une sixième liberté a "
        "un problème de conception que le design ne résout pas.")
saut()


# --------------------------------------------------------------------------- 02
surtitre("Les trois couches")
titre_chapitre(2, "Quand tout se discute, tout dérive.")
corps("Le noyau ne bouge presque jamais. Les densités bougent à l’arrivée d’un produit. "
      "L’écosystème bouge au fil des besoins. Trois rythmes dans un seul document, et "
      "c’est toujours le plus lent qui se fait emporter.")
etiquette("La répartition")
tableau(["Couche", "Rythme", "Ce qu’elle contient"], [
    ("Noyau", "Presque jamais. Le toucher est un événement, pas une correction",
     "Jetons de marque et jetons dérivés, typographie, huit composants, iconographie, voix, mode sombre"),
    ("Densités", "À l’ajout d’un produit, et seulement là",
     "Quatre profils d’espacement. Aucune couleur, aucune police, aucune taille de texte"),
    ("Écosystème", "Au fil des besoins inter-produits",
     "Menu de compte, sélecteur d’organisation, écran d’accès refusé, bandeau d’environnement"),
], [3.0, 5.0, 8.2])
etiquette("Le contrat de jetons")
corps("Aucune spécification de sprint ne cite une couleur brute. Elles citent toutes une "
      "variable. Une spécification qui écrit #2251FF est fausse ; une spécification qui "
      "écrit --action est juste. La différence n’est pas cosmétique : la première fige une "
      "valeur dans un document que personne ne relira, la seconde suit le système quand il "
      "corrige un contraste.")
etiquette("Ce que le système ne fera jamais")
corps("**Modifier la marque.** Les six jetons institutionnels sont recopiés et vérifiés, "
      "jamais amendés. Un produit qui redéfinit --marque-action dérive la marque sans que "
      "personne ne l’ait décidé.")
corps("**Régenter la vitrine.** Le site, les présentations et les documents commerciaux "
      "restent sous l’autorité de la charte mère. Ce système ne parle que des produits.")
encadre("**La couche écosystème est délibérément vide aujourd’hui.** Un composant "
        "inter-produits ne s’écrit qu’au moment où un deuxième produit le réclame. Écrit "
        "plus tôt, il fige une hypothèse au lieu d’un besoin.")
saut()


# --------------------------------------------------------------------------- 03
surtitre("Les jetons de marque")
titre_chapitre(3, "Six valeurs importées. Jamais redéfinies.")
corps("Un seul fichier du dépôt écrit une couleur de marque en clair, et il est généré. "
      "Tout le reste pointe vers lui. C’est la seule façon connue d’empêcher une marque de "
      "dériver : lui retirer les endroits où elle le pourrait.")
etiquette("Les six jetons")
tableau(["", "Valeur", "Jeton", "Source amont"], [
    ("#051C2C", "#051C2C", "--marque-encre", "--ai5d-ink"),
    ("#042A76", "#042A76", "--marque-navy", "--ai5d-navy"),
    ("#2251FF", "#2251FF", "--marque-action", "--ai5d-blue"),
    ("#1B44DB", "#1B44DB", "--marque-action-survol", "--ai5d-blue-hover"),
    ("#5B7BFF", "#5B7BFF", "--marque-action-clair", "--ai5d-blue-light"),
    ("#FFFFFF", "#FFFFFF", "--marque-blanc", "--ai5d-white"),
], [2.0, 2.6, 5.4, 6.2], nuancier=True)
etiquette("La copie est vérifiée, pas supposée")
corps("marque.css est produit par un générateur qui recopie ces six valeurs depuis la "
      "source institutionnelle et inscrit dans l’en-tête l’empreinte du fichier d’origine. "
      "Un test relit cette source à chaque exécution et échoue si la copie a divergé. "
      "Recopier sans vérifier revient à créer une seconde source de vérité qui se "
      "désynchronise en silence.")
code("/* Généré par _build/synchroniser-marque.mjs. Ne pas modifier à la main.\n"
     "   Source    : AI5D_Brand_2026/tokens.css\n"
     "   Empreinte : sha256(16) 03292c3bfad69b1b\n"
     "   Synchronisé le 2026-09-05 */")
etiquette("La répartition de la couleur")
corps("65 % surfaces · 25 % encre · 7 % action · 3 % sémantiques. Le bleu et le vert réunis "
      "ne dépassent jamais un dixième de la surface. Un bleu partout n’est plus un signal : "
      "il devient un décor, et l’action disparaît dedans.")
saut()


# --------------------------------------------------------------------------- 04
surtitre("Palette claire")
titre_chapitre(4, "Le blanc pur durcit. Le papier tiède accueille.")
corps("Vingt minutes de lecture sur du blanc pur, et l’œil se ferme. La marque "
      "institutionnelle n’a pas ce problème : on ne lit pas une page de présentation, on "
      "la parcourt. Ces jetons n’ont donc aucun équivalent en amont.")
etiquette("Surfaces")
tableau(["", "Valeur", "Jeton", "Rôle"], [
    ("#FAF7F2", "#FAF7F2", "--surface-1", "Fond de page. Tiède, jamais blanc pur"),
    ("#FFFFFF", "#FFFFFF", "--surface-2", "Cartes, panneaux, champs"),
    ("#FFFFFF", "#FFFFFF", "--surface-3", "Menus, dialogues, éléments flottants"),
    ("#F4EFE7", "#F4EFE7", "--surface-chaude", "Lectures longues, encadrés"),
    ("#E7E0D6", "#E7E0D6", "--bordure", "Filets courants"),
    ("#D5CCBE", "#D5CCBE", "--bordure-forte", "Filets appuyés, contours de contrôle"),
], [2.0, 2.6, 4.4, 7.2], nuancier=True)
etiquette("Texte et sémantiques, avec leur contraste le plus faible")
tableau(["", "Valeur", "Jeton", "Usage", "Ratio"], [
    ("#051C2C", "#051C2C", "--texte-fort", "Titres, chiffres, libellés", "15,18"),
    ("#2B3A45", "#2B3A45", "--texte", "Corps, texte courant", "10,23"),
    ("#616F78", "#616F78", "--texte-faible", "Aides, métadonnées, descriptions", "4,53"),
    ("#0E7C5A", "#0E7C5A", "--reussite", "Acquis, terminé, validé", "4,58"),
    ("#B45309", "#B45309", "--attention", "Adresse non confirmée, quota proche", "4,54"),
    ("#B42318", "#B42318", "--erreur", "Échec, validation refusée", "5,75"),
    ("#2251FF", "#2251FF", "--action", "Boutons, liens, repères", "4,96"),
], [1.8, 2.4, 3.8, 6.0, 2.2], nuancier=True)
corps("Ratio le plus faible des trois surfaces claires. Le seuil du texte courant est 4,5.",
      8.5, FAIBLE)
encadre("**Aucune information n’est portée par la seule couleur.** Réussite et erreur portent "
        "toujours aussi un mot ou une icône : près d’un homme sur douze ne distingue pas "
        "correctement le rouge du vert.")
saut()


# --------------------------------------------------------------------------- 05
surtitre("Palette sombre")
titre_chapitre(5, "Ici, la profondeur vient de la lumière. Pas de l’ombre.")
corps("Une partie des sessions se fait la nuit, sur téléphone, dans une pièce peu "
      "éclairée. Le mode sombre n’est donc pas une préférence : c’est le mode de référence "
      "d’une partie des écrans. Il se traite au premier écran, jamais au dernier. **Le "
      "choix du compte l’emporte sur la préférence du système.**")
etiquette("Surfaces et texte")
tableau(["", "Valeur", "Jeton", "Rôle"], [
    ("#0B1620", "#0B1620", "--surface-1", "Fond de page. Encre profonde, jamais noir pur"),
    ("#11212D", "#11212D", "--surface-2", "Cartes et panneaux"),
    ("#172C3B", "#172C3B", "--surface-3", "Menus, dialogues, éléments flottants"),
    ("#171F26", "#171F26", "--surface-chaude", "Lectures longues"),
    ("#6B88FF", "#6B88FF", "--action", "Le bleu s’éclaircit pour rester lisible"),
    ("#2FA37B", "#2FA37B", "--reussite", "Acquis, terminé"),
    ("#C9D4DC", "#C9D4DC", "--texte", "Titres à #F2F5F7, secondaire à #8D9AA5"),
], [2.0, 2.6, 4.4, 7.2], nuancier=True)
etiquette("Trois règles")
tableau(["Règle", "Ce qu’elle implique"], [
    ("Les ombres disparaissent",
     "Une ombre portée sur fond sombre ne se voit pas, et la simuler produit du gris sale. "
     "Les trois niveaux d’élévation valent none. Une surface plus claire est une surface "
     "plus proche : c’est là toute la hiérarchie."),
    ("Le bleu s’éclaircit",
     "C’est le seul jeton hérité dont la valeur change entre les deux thèmes. Le chapitre "
     "suivant explique pourquoi #5B7BFF ne suffisait pas."),
    ("Aucun scintillement",
     "Le thème est appliqué avant le premier rendu. Un flash blanc au chargement trahit une "
     "interface mal faite, et se remarque précisément la nuit."),
], [4.4, 11.8])
etiquette("Les trois états du thème")
tableau(["État", "Ce qui le déclenche", "Ce qui l’emporte"], [
    ("Préférence du système", "Aucun choix enregistré sur le compte",
     "La requête média du navigateur décide"),
    ("Choix explicite clair", "data-theme=\"light\" sur la racine",
     "Le compte, même si le système est en sombre"),
    ("Choix explicite sombre", "data-theme=\"dark\" sur la racine",
     "Le compte, même si le système est en clair"),
], [4.4, 6.0, 5.8])
corps("L’ordre compte : la règle du choix explicite est écrite après celle de la préférence "
      "système, sinon elle ne gagnerait jamais.", 8.5, FAIBLE)
saut()


# --------------------------------------------------------------------------- 06
surtitre("Les écarts mesurés")
titre_chapitre(6, "Quatre divergences. Trois bonnes décisions.")
corps("Quatre jetons sémantiques avaient divergé sans que personne ne l’ait décidé. Cela "
      "avait toutes les apparences d’une faute. Le calcul des contrastes a dit l’inverse.")
etiquette("Ce qui avait divergé")
tableau(["Jeton", "Marque mère", "Chartes dérivées"], [
    ("Réussite", "#1E874B", "#0E7C5A"),
    ("Erreur", "#D02B2B", "#B42318"),
    ("Avertissement", "#B7791F", "#B45309"),
    ("Texte secondaire", "#757575", "#6B7A85"),
], [5.0, 4.0, 7.2])
etiquette("Ce que la mesure a répondu")
tableau(["Valeur sur son fond", "Ratio", "Verdict"], [
    ("Réussite institutionnelle #1E874B sur blanc", "4,54", "Tient"),
    ("Réussite institutionnelle #1E874B sur papier tiède", "4,25✖", "Échoue✖"),
    ("Réussite applicative #0E7C5A sur papier tiède", "4,85", "Tient"),
    ("Avertissement institutionnel #B7791F sur blanc", "3,64✖", "Échoue✖"),
    ("Avertissement applicatif #B45309 sur papier tiède", "4,70", "Tient"),
    ("Erreur applicative #B42318 sur papier tiède", "6,15", "Tient"),
], [9.0, 2.6, 4.6], gras_col0=False)
corps("**Trois des quatre écarts étaient de bonnes décisions prises sans avoir été "
      "écrites.** Sur la surface tiède du registre applicatif, les valeurs "
      "institutionnelles lâchent. Le système les entérine donc, avec leur justification "
      "chiffrée, plutôt que de les corriger vers une valeur qui échoue.")
etiquette("Le quatrième était un vrai défaut, des deux côtés")
corps("Le texte secondaire hérité, #6B7A85, donnait 4,14 sur le papier tiède et 4,42 sur le "
      "blanc : il échouait partout. Une première correction à #66747E réglait le papier et "
      "le blanc mais restait à 4,20 sur la surface chaude, plus sombre que les deux autres. "
      "La valeur retenue, **#616F78**, tient les trois : 4,85 sur le fond de page, 5,18 sur "
      "le blanc, 4,53 sur la surface chaude. La teinte et la saturation sont conservées ; "
      "seule la luminosité baisse.")
etiquette("Et un défaut que seul le registre applicatif pouvait produire")
corps("La marque déclare --ai5d-blue-light #5B7BFF pour les liens sur fond encre ou navy. "
      "Sur l’encre, il tient : 4,73. Mais le registre applicatif introduit des surfaces que "
      "la marque n’a jamais eues, #11212D pour les cartes et #172C3B pour les menus, où il "
      "tombe à **4,47** et **3,92**. Un lien posé sur une carte, en mode sombre, échouait "
      "donc. Le jeton de marque n’a pas été touché : le noyau déclare à côté un jeton "
      "applicatif distinct, --action-sur-sombre à **#6B88FF**, qui tient sur les trois "
      "surfaces sombres et sur l’encre.")
etiquette("Les quatre défauts trouvés le jour où la garde a été écrite")
tableau(["Jeton", "Sur quelle surface", "Avant", "Après", "Nature"], [
    ("--texte-faible", "Les trois surfaces claires", "4,14✖", "4,53", "Défaut hérité de la charte"),
    ("--action, sombre", "Carte #11212D", "4,47✖", "5,15", "Surface inconnue de la marque"),
    ("--action, sombre", "Menu #172C3B", "3,92✖", "4,51", "Surface inconnue de la marque"),
    ("--reussite-fond, sombre", "Fond de pastille", "4,45✖", "4,63", "Valeur saisie sans mesure"),
], [4.0, 4.0, 1.8, 1.8, 4.6])
encadre("Ces quatre défauts n’ont pas été trouvés par relecture. Ils ont été trouvés par un "
        "test qui recalcule **quarante-quatre paires de contraste** à chaque exécution, en "
        "clair et en sombre, et échoue sous 4,5. Aucun n’était visible à l’œil, et trois "
        "étaient en production.")
saut()


# --------------------------------------------------------------------------- 07
surtitre("Typographie")
titre_chapitre(7, "Une serif pour signer. Une sans pour travailler.")
tableau(["Police", "Rôle", "Fichier"], [
    ("Fraunces 300–500",
     "Affichage et signature : titres d’écran, noms propres, en-têtes de courriel. Toutes "
     "les interfaces voisines composent en sans géométrique, et c’est elle qui rend un "
     "produit AI5D reconnaissable en une demi-seconde", "fraunces-variable.woff2"),
    ("Inter 400–700",
     "Interface et lecture : formulaires, navigation, journaux, mentions légales. Héritée "
     "de la maison mère, elle assure la continuité et supporte les sessions longues",
     "inter-variable.woff2"),
    ("JetBrains Mono 500",
     "Identification : codes à deux facteurs, empreintes de session, identifiants. Une "
     "chaîne qu’on recopie caractère par caractère se compose en chasse fixe",
     "jetbrains-mono-500.woff2"),
], [3.4, 8.4, 4.4])
etiquette("Hiérarchie")
tableau(["Niveau", "Police", "Taille", "Usage"], [
    ("Display", "Fraunces 300 · 400", "48 · 56 px", "Titre de page, écran d’accueil"),
    ("H1", "Fraunces 500", "30 px", "Titre d’écran"),
    ("H2", "Inter 600", "22 px", "Section"),
    ("H3", "Inter 600", "18 px", "Intitulé de carte"),
    ("Corps", "Inter 400", "16 px", "Texte courant"),
    ("Petit", "Inter 400", "14 px", "Aides, descriptions"),
    ("Légende", "Inter 500", "12 px", "Métadonnées, horodatage"),
    ("Overline", "Inter 600 · +0,06 em", "12 px", "Capitales uniquement ici"),
    ("Mono", "JetBrains Mono 500", "14 px", "Codes, empreintes, identifiants"),
], [3.0, 4.6, 2.6, 6.0])
encadre("**Trois règles.** Une seule graisse par bloc : la hiérarchie se fait par la taille, "
        "jamais par le gras. **Fraunces ne compose jamais un paragraphe.** Les capitales "
        "sont réservées aux overlines.")
corps("**134 Ko, en local.** Trois fichiers pour huit graisses : Inter et Fraunces sont "
      "variables, un seul fichier porte toute leur plage. Sous-ensemble latin. **Aucun "
      "appel à un service de polices tiers**, et un test le vérifie : une page "
      "d’authentification ne doit émettre aucune requête vers un tiers.")
saut()


# --------------------------------------------------------------------------- 08
surtitre("Les densités")
titre_chapitre(8, "Même ADN partout. Quatre respirations.")
corps("Une plateforme d’apprentissage et un atelier de données ne demandent pas le même "
      "air entre les éléments. C’est la seule chose qui change d’un produit AI5D à "
      "l’autre. Six variables d’espacement, et rien d’autre.")
etiquette("Les quatre profils")
tableau(["Variable", "Aéré · Académie", "Équilibré · Compte", "Modéré · Cercle", "Compact · Lab"], [
    ("Rythme de section", "64 px", "48 px", "40 px", "32 px"),
    ("Padding de carte", "32 px", "24 px", "20 px", "16 px"),
    ("Hauteur de contrôle", "48 px", "48 px", "44 px", "40 px"),
    ("Ligne de liste", "64 px", "56 px", "48 px", "40 px"),
    ("Interligne du corps", "1,6", "1,55", "1,5", "1,45"),
    ("Largeur de contenu", "1120 px", "1120 px", "1280 px", "pleine"),
], [4.4, 3.0, 3.2, 2.9, 2.7])
etiquette("Les deux règles")
tableau(["Règle", "Pourquoi elle ne se négocie pas"], [
    ("La densité change l’espace entre les choses, jamais la taille du texte",
     "Sans cette règle, le profil compact devient illisible en six mois : c’est la pente "
     "naturelle de tout profil dense, et la vigilance ne la corrige pas. Le fichier des "
     "profils ne contient donc ni couleur, ni famille, ni taille de police, et un test "
     "échoue s’il en apparaît une."),
    ("Le plancher tactile de 44 px prime sur les quatre profils",
     "Le Lab descend à 40 px sur un écran de bureau avec une souris. Jamais sur un "
     "téléphone : une cible de 40 px y produit des erreurs de saisie que la personne "
     "attribue à l’application, jamais à son doigt. La règle s’exprime une seule fois, en "
     "requête média, et n’est donc pas négociable écran par écran."),
], [5.4, 10.8])
code("@media (pointer: coarse) {\n"
     "  :root,\n"
     "  [data-densite] {\n"
     "    --hauteur-controle: max(var(--hauteur-controle), 44px);\n"
     "    --ligne-liste:      max(var(--ligne-liste), 44px);\n"
     "  }\n"
     "}")
corps("Le sélecteur générique [data-densite] couvre aussi les profils qu’on ajouterait plus "
      "tard et que personne n’aurait pensé à vérifier.", 8.5, FAIBLE)
etiquette("Ce qui ne varie jamais d’un profil à l’autre")
corps("Les couleurs. Les polices. Les tailles de texte. Les rayons. L’élévation. "
      "L’épaisseur des icônes. Les significations sémantiques. Le plancher tactile.")
corps("Un produit qui a besoin de changer l’un de ces éléments a un problème que la "
      "densité ne résout pas. C’est le sens de la première règle : elle ne décrit pas une "
      "préférence esthétique, elle empêche un produit de partir seul.")
etiquette("Ajouter un cinquième profil")
corps("C’est autorisé, quand un produit ne rentre dans aucun des quatre et que la "
      "démonstration est faite. Deux conditions de forme : le profil se déclare **avant** "
      "le bloc du plancher tactile, sinon il le remplacerait, et il ne touche qu’aux six "
      "variables d’espacement. Un test vérifie les deux.")
saut()


# --------------------------------------------------------------------------- 09
surtitre("Les huit composants")
titre_chapitre(9, "Un composant du noyau garantit ce qu’on oublie.")
corps("Ils ne sont pas dans le noyau parce qu’ils reviennent souvent. Ils y sont parce que "
      "chacun porte une règle qu’un développeur pressé enfreindrait sans le vouloir, et "
      "sans le voir. Aucun ne dépend d’un framework de style.")
etiquette("Ce que chacun garantit")
tableau(["Composant", "Sa garantie"], [
    ("Logotype", "Le « 5 » incliné à moins cinq degrés et bleu, dans toutes les variantes. "
                 "Par défaut les lettres suivent le texte fort, donc le logotype se retourne "
                 "tout seul avec le thème"),
    ("Bouton", "Trois variantes, trois tailles, hauteur pilotée par la densité, plancher "
               "tactile respecté, aria-busy pendant le chargement. Un seul bouton primaire par vue"),
    ("Champ", "Le libellé est toujours lié par htmlFor, l’aide et l’erreur par "
              "aria-describedby, et l’erreur n’est jamais portée par la seule couleur"),
    ("Carte", "Padding piloté par la densité. Cliquable, elle rend un <button>, jamais une "
              "<div> avec un gestionnaire de clic"),
    ("Bandeau", "Une icône et un texte. role=\"alert\" pour ce qui demande une réaction, "
                "role=\"status\" pour ce qui informe"),
    ("Pastille", "Un état compact, qui contient toujours du texte. Une pastille de couleur "
                 "sans mot ne dit rien"),
    ("Icone", "Lucide, contour, épaisseur 1,75. Décorative par défaut, accessible seulement "
              "si on lui donne un titre"),
    ("GabaritAuth", "Le gabarit d’authentification en deux colonnes, avec ses mesures. "
                    "Voir le chapitre 11"),
], [3.4, 12.8])
etiquette("Trois refus, et leur raison")
corps("**Pas de div cliquable.** Une carte actionnable qui n’est pas un bouton n’est ni "
      "atteignable au clavier ni annoncée comme actionnable.")
corps("**Pas d’icône seule.** Un cadenas ne dit pas quelle licence manque, ni à qui la "
      "demander. L’icône accompagne le mot, elle ne le remplace jamais dans une "
      "information d’état.")
corps("**Pas d’erreur en couleur seule.** Le champ pose aria-invalid, change la bordure et "
      "affiche un texte. Les trois, ensemble.")
encadre("**Pourquoi le gabarit d’authentification est dans le noyau.** Tout produit peut "
        "avoir à afficher un écran de session expirée, même si le portail Compte porte "
        "l’essentiel des flux. Le placer dans l’écosystème obligerait chaque produit à en "
        "dépendre pour un écran qu’il affiche une fois par mois.")
saut()


# --------------------------------------------------------------------------- 10 + 11
surtitre("Planches")
titre_chapitre(10, "Rien ne change. Sauf la lumière.")
encadre("**Ce chapitre n’existe pas dans cette version.** Les planches reproduisent des "
        "fragments d’interface : boutons dans leurs trois variantes, champs en état normal "
        "et en erreur, bandeaux et pastilles dans les quatre tons, logotype sur fond clair "
        "et sur fond encre, chacun rendu côte à côte dans les deux thèmes. Word ne sait "
        "pas les composer. **Voir le PDF, chapitre 10.**")
saut()

surtitre("Le gabarit d’authentification")
titre_chapitre(11, "Cette disposition n’est pas une proposition. Elle est mesurée.")
corps("Elle vient de l’Académie, où elle est en production. Chacune de ses valeurs a été "
      "relevée sur un écran réel, et chacune corrige un défaut constaté. Les raisons sont "
      "recopiées ici parce qu’un jour, quelqu’un voudra les changer.")
etiquette("La disposition")
tableau(["Palier", "Ce qui est affiché"], [
    ("À partir de 1024 px",
     "Deux colonnes. Le formulaire à gauche, dans 440 px. Un panneau d’encre à droite à "
     "45 %, plafonné à 560 px, qui porte le logotype et une phrase, et rien d’autre. Ni "
     "photo, ni illustration, ni filet, ni forme animée : cette sobriété est le signal de "
     "sérieux qui distingue un produit AI5D des plateformes grand public."),
    ("Sous 1024 px",
     "Colonne unique, logotype centré au-dessus à 20 px, marge de 16 px, aucune image "
     "lourde : l’écran doit s’afficher en moins de deux secondes et demie sur une 4G "
     "irrégulière. Le contenu court se centre dans la hauteur restante ; un formulaire "
     "long pousse simplement le conteneur."),
], [4.4, 11.8])
etiquette("Trois mesures, et ce qu’elles ont coûté")
tableau(["Mesure", "Ce qu’elle corrige"], [
    ("La bascule est à 1024 px, et non 768",
     "À 768 px, le panneau prend 45 %, soit 345 px, et il reste 423 px pour un formulaire "
     "annoncé à 440 px plus ses marges : l’écran de réinitialisation débordait de 14 px. "
     "À 1024 px, la colonne du formulaire dispose de 563 px et tout respire."),
    ("min-width: 0 sur la colonne du formulaire",
     "Ce n’est pas décoratif. Sans cette déclaration, une colonne flexible refuse de "
     "descendre sous la largeur intrinsèque de son contenu et pousse le panneau hors de "
     "l’écran au lieu de se réduire. C’est la cause exacte du débordement mesuré."),
    ("Le logotype mobile est à 20 px, et non 18",
     "À 18 px, le verrouillage complet mesurait 19 px de haut, sous le plancher de 28 px "
     "de la charte : il se voyait mal. On ne monte pas plus haut non plus, car à 30 px la "
     "signature réclame 250 px de large et se casse en deux lignes sur un écran de 320 px."),
], [5.0, 11.2])
etiquette("Les cinq écrans que le gabarit porte")
tableau(["Écran", "Ce qui change dans la colonne de gauche"], [
    ("Connexion", "Deux champs et le bouton primaire. Le lien d’oubli en variante discrète, sous le bouton"),
    ("Inscription", "Le formulaire le plus long des cinq : c’est lui qui pousse le conteneur au lieu de se centrer"),
    ("Oubli du mot de passe", "Un seul champ. Contenu court, donc centré dans la hauteur restante"),
    ("Réinitialisation", "Deux champs. C’est cet écran qui débordait de 14 px à 768 px"),
    ("Vérification d’adresse", "Aucun champ. Un titre, une phrase, et l’action de renvoi en variante discrète"),
], [4.4, 11.8])
encadre("**Ce que le panneau ne porte plus.** Il a porté trois preuves, avec des chiffres et "
        "des garanties. Elles ont été retirées : elles promettaient un contenu qu’un écran "
        "de connexion n’a pas à vendre. On y arrive déjà décidé.")
saut()


# --------------------------------------------------------------------------- 12
surtitre("Mobile d’abord")
titre_chapitre(12, "Le téléphone n’est pas un palier. C’est le départ.")
corps("Un écran dessiné pour 1440 px puis rétréci garde les décisions du grand écran : une "
      "grille qui s’empile mal, une navigation qui déborde, des cibles calibrées pour un "
      "curseur. Un écran dessiné pour 390 px puis élargi n’a que de la place en plus. Les "
      "espaces d’administration sont la seule exception.")
etiquette("Les paliers")
tableau(["Constante", "Valeur", "Nature", "Ce qui s’y passe"], [
    ("--plancher", "320 px", "contrainte", "Rien. Aucune mise en page n’a le droit d’y casser"),
    ("--reference-mobile", "390 px", "contrainte", "Rien. C’est la largeur sur laquelle on dessine"),
    ("--compact", "640 px", "palier", "La marge de page passe de 16 à 24 px"),
    ("--tablette", "768 px", "palier", "La barre d’onglets disparaît, la navigation remonte"),
    ("--bureau", "1024 px", "palier", "Deux colonnes deviennent possibles"),
    ("--large", "1280 px", "palier", "Le contenu est plafonné et centré"),
], [4.4, 2.4, 2.8, 6.6])
corps("**Deux des six ne sont pas des paliers.** Rien ne s’y déclenche. Les distinguer évite "
      "qu’on écrive un jour @media (min-width: 320px), qui ne voudrait rien dire.")
etiquette("Pourquoi 640 et non 480")
corps("Le formulaire du chapitre 11 vaut 440 px, et deux marges de 32 px en ajoutent 64. Il "
      "faut donc **au moins 504 px** avant d’élargir les marges. 640 est la valeur ronde "
      "immédiatement au-dessus. Le palier bureau retrouve indépendamment le 1024 px du même "
      "chapitre, et un test croise les deux valeurs.")
etiquette("Les sept règles")
tableau(["Règle", "Ce qu’elle empêche"], [
    ("Le pouce d’abord",
     "L’action principale vit dans le tiers bas. Sur un téléphone de six pouces tenu d’une "
     "main, le coin haut droit demande de changer la prise"),
    ("La zone sûre se réserve",
     "Sans env(safe-area-inset-*), une barre basse passe sous la barre de gestes d’un "
     "iPhone : visible, et inatteignable"),
    ("dvh, jamais vh",
     "La barre d’URL entre et sort du cadre en défilant. Un écran calé sur 100vh se fait "
     "couper au chargement"),
    ("Aucune largeur figée au-delà de 320 px",
     "max-width est la solution, width: 440px est le problème : il ne descend pas quand "
     "l’écran descend"),
    ("Le survol n’apprend rien",
     "Au doigt, :hover se déclenche après la pression et reste collé. Ce qui n’apparaît "
     "qu’au survol n’existe pas"),
    ("Le plancher tactile prime",
     "44 px, déjà garanti par les profils de densité. Les paliers ne le renégocient pas"),
    ("Une colonne sous 640 px",
     "Deux colonnes sur 390 px donnent des cibles de 150 px de large. On empile"),
], [5.0, 11.2])
etiquette("La coquille d’application")
corps("Trois zones : en-tête collant de 56 px, contenu défilant, barre d’onglets de 56 px. "
      "**La mesure qui décide de tout est la réserve basse** : le contenu réserve sous lui la "
      "hauteur de la barre plus la zone sûre. Sans elle, le dernier élément se glisse sous la "
      "barre. Le défaut ne se voit pas sur une page courte, et apparaît le jour où quelqu’un "
      "ajoute une ligne.")
corps("**Trois à cinq onglets.** Sous trois, deux liens tiennent dans l’en-tête. Au-delà de "
      "cinq, chaque cible descend sous 70 px sur un téléphone de 390 px. Le sixième devient "
      "« Plus ». **L’onglet actif ne se signale pas par la seule couleur** : il porte "
      "aria-current, prend la couleur d’action, et passe sa graisse à semi-grasse.")
etiquette("Ce qu’on refuse au motif de référence")
tableau(["Élément du motif", "Décision", "Raison"], [
    ("Une couleur par section", "Refusé✖",
     "La charte mère écrit que la différenciation se fait par le nom, jamais par la couleur"),
    ("Fond de carte teinté", "Refusé✖",
     "Vert et jaune sont des jetons sémantiques. Décoratifs, « réussite » ne veut plus rien "
     "dire ailleurs"),
    ("Boutons entièrement arrondis", "Refusé✖",
     "Le rayon de bouton vaut 10 px. Le rayon plein est réservé aux pastilles et aux avatars"),
    ("Quatre onglets dont « Plus »", "Repris, et borné",
     "Trois à cinq, et pas davantage"),
], [5.0, 3.0, 8.2])
encadre("**Ce que les gardes ne peuvent pas prouver.** Elles lisent des fichiers, pas des "
        "écrans rendus. Une largeur figée se voit ; un tableau illisible sur 390 px, non. Cela "
        "se vérifie à l’écran, et nulle part ailleurs.")
saut()


# --------------------------------------------------------------------------- 13
surtitre("Iconographie et mouvement")
titre_chapitre(13, "Un quart de pixel qui se justifie.")
etiquette("L’iconographie")
corps("**Lucide, style contour, épaisseur 1,75 px**, contre 1,5 dans le registre "
      "institutionnel. Les 0,25 px se justifient sur un écran de téléphone, où un trait de "
      "1,5 disparaît. Jamais d’icône remplie. Tailles 16, 20, 24, 32 et 72.")
corps("Par défaut, une icône est **décorative** et masquée aux lecteurs d’écran. Ce n’est "
      "pas de la négligence : c’est la règle. Quand l’icône porte réellement "
      "l’information, et cela devrait être rare, on lui donne un titre.")
etiquette("La géométrie et l’élévation")
tableau(["Jeton", "Valeur", "Usage"], [
    ("--rayon-sm", "4 px", "Boutons, champs, pastilles"),
    ("--rayon-md", "10 px", "Cartes, panneaux"),
    ("--rayon-lg", "16 px", "Dialogues, grands blocs"),
    ("--rayon-plein", "999 px", "Pastilles, avatars"),
    ("--elevation-1", "0 1px 2px rgba(5,28,44,.06)", "Repos, listes"),
    ("--elevation-2", "0 4px 12px rgba(5,28,44,.08)", "Cartes actionnables"),
    ("--elevation-3", "0 12px 32px rgba(5,28,44,.12)", "Dialogues, menus"),
], [4.0, 6.2, 6.0])
corps("Les trois élévations valent none en mode sombre. Voir le chapitre 05.", 8.5, FAIBLE)
etiquette("Le mouvement")
tableau(["Durée", "Valeur", "Usage"], [
    ("Courte", "150 ms", "Survol, pression, apparition d’une aide"),
    ("Moyenne", "250 ms", "Ouverture d’un menu, bascule de panneau"),
    ("Longue", "800 ms", "Confirmation qui engage la sécurité du compte, jamais un ornement"),
], [3.0, 3.0, 10.2])
corps("Courbe d’entrée cubic-bezier(0.16, 1, 0.3, 1), courbe de sortie "
      "cubic-bezier(0.4, 0, 1, 1). Sous prefers-reduced-motion, la durée longue tombe à "
      "zéro et les deux autres à 100 ms : il ne reste que des transitions d’opacité "
      "imperceptibles.", 8.5, FAIBLE)
saut()


# --------------------------------------------------------------------------- 13
surtitre("La voix")
titre_chapitre(14, "Un verrou dit toujours ce qui le lève.")
corps("La vitrine parle quand tout va bien. Le produit parle au moment du refus, de la "
      "panne, de la suppression de compte. C’est là que le ton se joue, et c’est là qu’il "
      "est le plus souvent raté.")
etiquette("Les règles propres à l’applicatif")
corps("**Vouvoiement partout**, y compris dans les erreurs et les courriels. **Aucun tiret "
      "cadratin.** Ni emoji, ni point d’exclamation, nulle part. **Nommer la "
      "conséquence** : un verrou dit ce qui le lève, un avertissement dit ce qu’on perd. "
      "**Pas d’anglicisme** quand le français existe. **Assumer l’absence** plutôt "
      "qu’inventer une réponse.")
etiquette("Formulations de référence")
tableau(["Situation", "Formulation", "La règle derrière"], [
    ("Échec de connexion", "« Adresse ou mot de passe incorrect. »",
     "Jamais lequel des deux : le dire révèle quelles adresses existent"),
    ("Mot de passe oublié",
     "« Si un compte existe pour cette adresse, vous recevrez un lien dans quelques minutes. »",
     "La réponse est identique que l’adresse existe ou non"),
    ("Compte verrouillé", "« Trop de tentatives. Réessayez dans cinq minutes. »",
     "On dit toujours quand le verrou se lève"),
    ("Session expirée",
     "« Votre session a expiré. Reconnectez-vous pour reprendre où vous en étiez. »",
     "On promet la reprise, sinon la personne craint d’avoir tout perdu"),
    ("Accès refusé, organisation",
     "« Votre organisation n’a pas d’accès à ce produit. Demandez-le à son administrateur. »",
     "On nomme qui peut lever le blocage"),
    ("Accès expiré", "« Votre accès a pris fin le 4 septembre. »",
     "On donne la date, jamais une durée relative"),
    ("Avant une suppression",
     "« Votre compte sera supprimé dans trente jours, avec vos accès à tous les produits "
     "AI5D. D’ici là, une simple connexion annule la demande. »",
     "On nomme la conséquence et la sortie"),
    ("Après une réinitialisation",
     "« Votre mot de passe est modifié. Toutes vos autres sessions ont été fermées. »",
     "On dit l’effet de bord : il rassure au lieu d’inquiéter"),
    ("Information indisponible",
     "« Nous ne pouvons pas afficher cette information pour le moment. »",
     "Assumer l’absence plutôt qu’afficher un tiret"),
], [3.6, 7.0, 5.6])
corps("Ce répertoire existe pour qu’un même refus ne soit pas dit de quatre façons dans "
      "quatre produits. Un produit y puise plutôt que de réinventer sa formulation.",
      8.5, FAIBLE)
saut()


# --------------------------------------------------------------------------- 14
surtitre("Les gardes")
titre_chapitre(15, "Une règle qu’aucune machine ne vérifie sera enfreinte.")
corps("Quatre jetons ont divergé sans décision. Aucun n’était visible à l’œil. Trois "
      "étaient en production. La discipline humaine a échoué une fois ; elle échouera "
      "encore. Le système livre donc des gardes que chaque projet branche dans son "
      "intégration continue.")
etiquette("Les trois gardes distribuées")
tableau(["Garde", "Ce qu’elle refuse, et pourquoi"], [
    ("aucune-couleur-en-dur",
     "Toute valeur hexadécimale ou fonction de couleur dans un fichier qui n’est pas un "
     "fichier de définition de jetons. Un écran qui décide une couleur dans son coin est "
     "un écran qui dérivera."),
    ("aucun-jeton-de-marque-redefini",
     "Toute ligne qui redéfinit l’un des six jetons de marque. C’est exactement le "
     "mécanisme qui a produit les écarts constatés."),
    ("cible-tactile-minimale",
     "L’absence de la requête média qui relève la hauteur de contrôle à 44 px sur "
     "pointeur grossier. Une règle unique protège tous les écrans à la fois ; sa "
     "disparition les découvre tous d’un coup."),
], [5.4, 10.8])
corps("Ce sont des **fonctions pures**. Elles parcourent une arborescence et rendent une "
      "liste d’infractions. Elles ne lèvent pas d’exception, n’affichent rien et ne "
      "décident rien : c’est le projet consommateur qui décide de faire échouer sa "
      "construction sur une liste non vide. Une garde qui déciderait à la place du projet "
      "serait désactivée le jour où elle gênerait.")
code("import { verifierAucuneCouleurEnDur, decrire } from '@ai5d/design-system/gardes';\n"
     "\n"
     "const infractions = verifierAucuneCouleurEnDur('apps/compte', {\n"
     "  exceptions: ['app/globals.css'],\n"
     "});\n"
     "\n"
     "if (infractions.length > 0) {\n"
     "  throw new Error(decrire(infractions));\n"
     "}")
etiquette("Les gardes internes au système")
tableau(["Garde", "Ce qu’elle vérifie"], [
    ("Contraste", "Recalcule 44 paires à chaque exécution, en clair et en sombre, et échoue "
                  "sous 4,5. C’est ce test qui aurait attrapé les quatre défauts dès le premier jour."),
    ("Fidélité de la marque", "Relit la source institutionnelle et échoue si la copie des "
                                "six jetons a divergé."),
    ("Pureté des densités", "Échoue si une couleur, une famille ou une taille de police "
                              "apparaît dans le fichier des profils, et vérifie que le plancher "
                              "tactile est déclaré en dernier."),
    ("Autonomie des polices", "Échoue si un appel vers un service de polices tiers apparaît "
                              "quelque part."),
], [4.6, 11.6])
encadre("**Ce qu’aucune garde ne peut faire.** Elles lisent des fichiers, pas des écrans "
        "rendus. Un contraste juste entre deux jetons n’empêche pas un écran de poser du "
        "texte faible sur une image. La relecture visuelle reste nécessaire ; les gardes "
        "lui retirent seulement le travail qu’elle fait mal.")
saut()


# --------------------------------------------------------------------------- 15
surtitre("Adoption")
titre_chapitre(16, "Un import, un attribut. C’est toute l’adoption.")
etiquette("Les deux gestes")
code("/* apps/<produit>/app/globals.css */\n@import '@ai5d/design-system/preset';")
code("<!-- apps/<produit>/app/layout.tsx -->\n<html lang=\"fr\" data-densite=\"equilibre\">")
corps("Le préréglage tire derrière lui les polices, les jetons de marque, les jetons du "
      "noyau et les quatre profils de densité. Les composants lisent la hauteur de contrôle "
      "et le padding de carte sans jamais savoir quel profil est actif.", 8.5, FAIBLE)
etiquette("Choisir son profil")
tableau(["Produit", "Profil", "Pourquoi"], [
    ("Académie", "aere", "Lecture, apprentissage, respiration"),
    ("Compte", "equilibre", "Gestion, sécurité, paramètres"),
    ("Le Cercle", "modere", "Communauté, interactions, flux"),
    ("Le Lab", "compact", "Données, workflows, outils, expérimentation"),
], [3.6, 3.4, 9.2])
etiquette("Le versionnement")
encadre("**Un changement de valeur de jeton est une version majeure.** Il modifie le rendu "
        "de tous les produits qui consomment le système, y compris ceux dont personne ne "
        "s’occupe ce mois-ci. Une correction de contraste se lit donc dans le numéro de "
        "version, avant même d’ouvrir le journal des changements.")
etiquette("Ce qui reste à écrire")
tableau(["Chantier", "État"], [
    ("Les composants inter-produits",
     "Menu de compte, sélecteur d’organisation, écran d’accès refusé, bandeau "
     "d’environnement. Ils attendent leur deuxième consommateur : écrits plus tôt, ils "
     "figeraient une hypothèse."),
    ("Les logotypes vectorisés",
     "Les fichiers actuels dépendent d’une police pour composer le mot-symbole. En "
     "courriel et en export hors ligne, le nom AI5D se rend donc dans une police de "
     "substitution."),
    ("La migration de l’Académie",
     "Le produit consomme aujourd’hui sa propre charte. Il rejoindra le système quand le "
     "portail Compte l’aura éprouvé."),
], [5.0, 11.2])
saut()


# --------------------------------------------------------------------------- fin
p = para(90, 6, 1.2, WD_ALIGN_PARAGRAPH.CENTER)
police(p.add_run("Un seul ADN."), SERIF, 22, ENCRE)
p = para(0, 14, 1.2, WD_ALIGN_PARAGRAPH.CENTER)
police(p.add_run("quatre respirations."), SERIF, 22, ACTION)
p = para(0, 30, 1.5, WD_ALIGN_PARAGRAPH.CENTER)
police(p.add_run("Source de vérité du registre applicatif AI5D. Toute évolution passe par une "
                 "nouvelle édition de ce document et une mise à jour de noyau/jetons.css."),
       SANS, 9, FAIBLE)

doc.save(os.path.abspath(SORTIE))
print("ecrit :", os.path.abspath(SORTIE))
